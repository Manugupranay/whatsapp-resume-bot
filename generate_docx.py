"""
Generates a professional resume DOCX from a tailored resume dict.
Uses python-docx for clean, ATS-friendly formatting.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io


def _add_horizontal_rule(doc):
    """Adds a thin horizontal line (border on paragraph)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E4057")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _section_heading(doc, text):
    """Adds a bold, coloured section heading with underline rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
    _add_horizontal_rule(doc)


def _bullet(doc, text, indent_level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def generate_docx(resume: dict) -> bytes:
    """
    Takes a tailored resume dict and returns DOCX file as bytes.
    """
    doc = Document()

    # --- Page margins ---
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # --- Set default font ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── NAME ──
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(resume["name"])
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    # ── CONTACT LINE ──
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(4)
    contact_run = contact_para.add_run(
        f"Email: {resume['email']}  |  Phone: {resume['phone']}"
    )
    contact_run.font.size = Pt(10)

    # ── PROFESSIONAL SUMMARY ──
    _section_heading(doc, "Professional Summary")
    summary_para = doc.add_paragraph()
    summary_para.paragraph_format.space_after = Pt(4)
    run = summary_para.add_run(resume["summary"])
    run.font.size = Pt(10)

    # ── TECHNICAL SKILLS ──
    _section_heading(doc, "Technical Skills")
    for category, skills_text in resume["skills"].items():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        bold_run = p.add_run(f"{category}: ")
        bold_run.bold = True
        bold_run.font.size = Pt(10)
        normal_run = p.add_run(skills_text)
        normal_run.font.size = Pt(10)

    # ── PROFESSIONAL EXPERIENCE ──
    _section_heading(doc, "Professional Experience")
    for job in resume["experience"]:
        # Company + dates line
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        company_run = p.add_run(f"{job['company']} – {job['location']}")
        company_run.bold = True
        company_run.font.size = Pt(10)
        # Right-align dates using tab stop
        tab_stop_pos = int(Inches(7.0).pt * 20)  # in twips
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(7.0), WD_ALIGN_PARAGRAPH.RIGHT
        )
        dates_run = p.add_run(f"\t{job['dates']}")
        dates_run.font.size = Pt(10)

        # Title line
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_before = Pt(0)
        title_p.paragraph_format.space_after = Pt(2)
        title_run = title_p.add_run(job["title"])
        title_run.bold = True
        title_run.italic = True
        title_run.font.size = Pt(10)
        title_run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

        # Bullets
        for bullet_text in job["bullets"]:
            _bullet(doc, bullet_text)

    # ── EDUCATION ──
    _section_heading(doc, "Education")
    for edu in resume["education"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        deg_run = p.add_run(edu["degree"])
        deg_run.bold = True
        deg_run.font.size = Pt(10)
        school_run = p.add_run(f" – {edu['school']}")
        school_run.font.size = Pt(10)

    # ── CERTIFICATIONS ──
    _section_heading(doc, "Certifications")
    for cert in resume["certifications"]:
        _bullet(doc, cert)

    # ── Serialize to bytes ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
